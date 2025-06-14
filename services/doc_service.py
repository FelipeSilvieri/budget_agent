from __future__ import annotations

import os
from datetime import datetime
from typing import List, Optional

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from pydantic import BaseModel

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "..", "static")
if not os.path.exists(OUTPUT_DIR):
    os.makedirs(OUTPUT_DIR)


class BudgetItem(BaseModel):
    descricao: str
    qtde: Optional[int] = 0
    preco_unitario: Optional[float] = 0.0


class BudgetRequest(BaseModel):
    cliente: Optional[str] = None
    responsavel: Optional[str] = None
    proposta_comercial: Optional[str] = None
    email: Optional[str] = None
    data: Optional[str] = None
    assunto: Optional[str] = None
    itens: List[BudgetItem]


def generate_document(req: BudgetRequest) -> str:
    """Gera um arquivo DOCX de orçamento e retorna o path relativo."""
    doc = Document()
    logo_path = os.path.join(os.path.dirname(__file__), "..", "logo_portal_center.png")
    if os.path.exists(logo_path):
        doc.add_picture(logo_path, width=Inches(5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    h = doc.add_heading("Orçamento", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    tbl1 = doc.add_table(rows=6, cols=2)
    tbl1.style = "Table Grid"
    labels = ["Cliente", "A/C", "Proposta", "E-mail", "Data", "Assunto"]
    vals = [
        req.cliente or "",
        req.responsavel or "",
        req.proposta_comercial or "",
        req.email or "",
        req.data or datetime.now().strftime("%d/%m/%Y"),
        req.assunto or "",
    ]
    for i, label in enumerate(labels):
        tbl1.columns[0].cells[i].text = label
        tbl1.columns[1].cells[i].text = vals[i]

    doc.add_paragraph("")

    for item in req.itens:
        if item.preco_unitario is None:
            item.preco_unitario = 0.0
        if item.qtde is None:
            item.qtde = 0

    df = pd.DataFrame(
        [
            {
                "ID Item": idx + 1,
                "Qtde": item.qtde,
                "R$ Unit": item.preco_unitario,
                "R$ Total": float(item.qtde or 0) * float(item.preco_unitario or 0.0),
                "Descrição": item.descricao,
            }
            for idx, item in enumerate(req.itens)
        ]
    )

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    for ci, col in enumerate(df.columns):
        hdr[ci].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)

    total = df["R$ Total"].sum()
    row = table.add_row()
    cell = row.cells[0]
    cell.merge(row.cells[-1])
    cell.text = f"Valor Total: R$ {total:.2f}"
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    filename = f"orcamento_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    path = os.path.join(OUTPUT_DIR, filename)
    doc.save(path)
    return f"/static/{filename}"
